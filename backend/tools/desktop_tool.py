"""
Desktop Control Tool — Agent GUI Automation
Gives agents full desktop control: mouse, keyboard, file management,
document editing, and web browsing via Playwright.

Dependencies:
    pip install pyautogui pillow playwright python-docx openpyxl pygetwindow
    playwright install chromium

Supported operations:
  Mouse    : move, click, double_click, right_click, drag, scroll
  Keyboard : type_text, press_key, hotkey, key_down, key_up
  Screen   : screenshot, find_on_screen, get_screen_size, get_mouse_position
  Files    : open_file, save_file, close_file, create_file, delete_file,
             copy_file, move_file, list_directory
  Documents: read_document, edit_document, save_document, create_document
             (supports .txt .docx .xlsx .csv .json .md .pdf)
  Browser  : browse_to, browser_click, browser_type, browser_screenshot,
             browser_get_text, browser_scroll, browser_back, browser_forward,
             browser_close, browser_find_element, browser_execute_js,
             browser_fill_form, browser_download

Authorized tiers: 0xxxx (Head), 1xxxx (Council), 2xxxx (Lead)
"""

import os
import re
import csv
import json
import time
import shutil
import asyncio
import logging
import platform
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ── Optional dependency guards ─────────────────────────────────────────────────

try:
    import pyautogui
    pyautogui.FAILSAFE = True   # Move mouse to top-left corner to abort
    pyautogui.PAUSE = 0.05      # Small delay between actions for stability
    PYAUTOGUI_AVAILABLE = True
except ImportError:
    PYAUTOGUI_AVAILABLE = False
    logger.warning("pyautogui not installed — mouse/keyboard control unavailable")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — .docx editing unavailable")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl not installed — .xlsx editing unavailable")

try:
    from playwright.async_api import async_playwright, Browser, Page
    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False
    logger.warning("playwright not installed — browser automation unavailable")


# ── Helpers ────────────────────────────────────────────────────────────────────

def _ts() -> str:
    return datetime.utcnow().isoformat()

def _require(flag: bool, name: str) -> Optional[Dict]:
    if not flag:
        return {"status": "error", "error": f"{name} not installed. Run: pip install {name.lower()}"}
    return None


# ══════════════════════════════════════════════════════════════════════════════
# MOUSE & KEYBOARD
# ══════════════════════════════════════════════════════════════════════════════

class MouseKeyboardTool:
    """
    Full mouse and keyboard control via pyautogui.
    Works on Windows, macOS, and Linux (requires display on Linux — use Xvfb if headless).
    """

    # ── Mouse ──────────────────────────────────────────────────────────────────

    def move(self, x: int, y: int, duration: float = 0.2) -> Dict[str, Any]:
        """Move mouse to absolute screen coordinates."""
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        pyautogui.moveTo(x, y, duration=duration)
        return {"status": "success", "action": "move", "x": x, "y": y}

    def click(self, x: int, y: int, button: str = "left", clicks: int = 1) -> Dict[str, Any]:
        """Click at coordinates. button: left | right | middle."""
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        pyautogui.click(x, y, button=button, clicks=clicks, interval=0.1)
        return {"status": "success", "action": "click", "x": x, "y": y, "button": button}

    def double_click(self, x: int, y: int) -> Dict[str, Any]:
        """Double-click at coordinates."""
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        pyautogui.doubleClick(x, y)
        return {"status": "success", "action": "double_click", "x": x, "y": y}

    def right_click(self, x: int, y: int) -> Dict[str, Any]:
        """Right-click at coordinates."""
        return self.click(x, y, button="right")

    def drag(self, from_x: int, from_y: int, to_x: int, to_y: int, duration: float = 0.5) -> Dict[str, Any]:
        """Click and drag from one position to another."""
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        pyautogui.moveTo(from_x, from_y)
        pyautogui.dragTo(to_x, to_y, duration=duration, button="left")
        return {"status": "success", "action": "drag",
                "from": [from_x, from_y], "to": [to_x, to_y]}

    def scroll(self, x: int, y: int, clicks: int = 3, direction: str = "down") -> Dict[str, Any]:
        """Scroll at position. clicks=number of scroll steps, direction: up|down|left|right."""
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        amount = clicks if direction == "up" else -clicks
        if direction in ("left", "right"):
            h_amount = clicks if direction == "right" else -clicks
            pyautogui.hscroll(h_amount)
        else:
            pyautogui.scroll(amount, x=x, y=y)
        return {"status": "success", "action": "scroll", "direction": direction, "clicks": clicks}

    def get_position(self) -> Dict[str, Any]:
        """Get current mouse cursor position."""
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        pos = pyautogui.position()
        return {"status": "success", "x": pos.x, "y": pos.y}

    # ── Keyboard ───────────────────────────────────────────────────────────────

    def type_text(self, text: str, interval: float = 0.02) -> Dict[str, Any]:
        """Type a string of text at the current cursor position."""
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        pyautogui.typewrite(text, interval=interval)
        return {"status": "success", "action": "type_text", "length": len(text)}

    def press_key(self, key: str) -> Dict[str, Any]:
        """
        Press a single key. Supports all pyautogui key names:
        enter, tab, backspace, delete, escape, space, up, down, left, right,
        home, end, pageup, pagedown, f1-f12, ctrl, alt, shift, win, cmd, etc.
        """
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        pyautogui.press(key)
        return {"status": "success", "action": "press_key", "key": key}

    def hotkey(self, keys: List[str]) -> Dict[str, Any]:
        """
        Press a key combination simultaneously.
        Examples: ['ctrl','c'], ['ctrl','alt','delete'], ['cmd','space'] on macOS
        """
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        pyautogui.hotkey(*keys)
        return {"status": "success", "action": "hotkey", "keys": keys}

    def key_down(self, key: str) -> Dict[str, Any]:
        """Hold a key down (use key_up to release)."""
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        pyautogui.keyDown(key)
        return {"status": "success", "action": "key_down", "key": key}

    def key_up(self, key: str) -> Dict[str, Any]:
        """Release a held key."""
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        pyautogui.keyUp(key)
        return {"status": "success", "action": "key_up", "key": key}

    # ── Screen ─────────────────────────────────────────────────────────────────

    def screenshot(self, save_path: str = "/tmp/desktop_screenshot.png") -> Dict[str, Any]:
        """Take a screenshot of the entire desktop."""
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        img = pyautogui.screenshot()
        img.save(save_path)
        return {"status": "success", "action": "screenshot", "path": save_path,
                "width": img.width, "height": img.height}

    def get_screen_size(self) -> Dict[str, Any]:
        """Return the screen resolution."""
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        w, h = pyautogui.size()
        return {"status": "success", "width": w, "height": h}

    def find_on_screen(self, image_path: str, confidence: float = 0.9) -> Dict[str, Any]:
        """
        Find a reference image on screen and return its centre coordinates.
        Requires opencv-python for confidence parameter.
        """
        err = _require(PYAUTOGUI_AVAILABLE, "pyautogui")
        if err: return err
        try:
            location = pyautogui.locateOnScreen(image_path, confidence=confidence)
            if location is None:
                return {"status": "error", "error": "Image not found on screen"}
            centre = pyautogui.center(location)
            return {"status": "success", "found": True, "x": centre.x, "y": centre.y,
                    "region": {"left": location.left, "top": location.top,
                               "width": location.width, "height": location.height}}
        except Exception as e:
            return {"status": "error", "error": str(e)}


# ══════════════════════════════════════════════════════════════════════════════
# FILE MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════

class FileManagementTool:
    """
    File system operations: create, read, edit, copy, move, delete, list.
    Paths on the host are accessed via /host mount inside Docker.
    """

    HOST_PREFIX = "/host"

    def _host_path(self, path: str) -> str:
        if path.startswith(self.HOST_PREFIX):
            return path
        if os.path.exists(self.HOST_PREFIX) and not path.startswith("/tmp"):
            return os.path.join(self.HOST_PREFIX, path.lstrip("/"))
        return path

    def open_file(self, filepath: str) -> Dict[str, Any]:
        """Open a file with the OS default application."""
        filepath = self._host_path(filepath)
        system = platform.system().lower()
        try:
            if system == "darwin":
                subprocess.Popen(["open", filepath])
            elif system == "windows":
                os.startfile(filepath)
            else:
                subprocess.Popen(["xdg-open", filepath])
            return {"status": "success", "action": "open_file", "path": filepath}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    def create_file(self, filepath: str, content: str = "") -> Dict[str, Any]:
        """Create a new file with optional initial content."""
        filepath = self._host_path(filepath)
        try:
            Path(filepath).parent.mkdir(parents=True, exist_ok=True)
            Path(filepath).write_text(content, encoding="utf-8")
            return {"status": "success", "action": "create_file",
                    "path": filepath, "bytes": len(content)}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    def read_file(self, filepath: str, offset: int = 0, limit: int = 500) -> Dict[str, Any]:
        """Read file contents with optional line offset and limit."""
        filepath = self._host_path(filepath)
        try:
            lines  = Path(filepath).read_text(encoding="utf-8", errors="replace").splitlines()
            slice_ = lines[offset: offset + limit]
            return {"status": "success", "path": filepath, "content": "\n".join(slice_),
                    "total_lines": len(lines), "returned_lines": len(slice_), "offset": offset}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    def save_file(self, filepath: str, content: str, backup: bool = True) -> Dict[str, Any]:
        """Write content to file, optionally creating a .bak backup first."""
        filepath = self._host_path(filepath)
        try:
            p = Path(filepath)
            if backup and p.exists():
                shutil.copy2(filepath, f"{filepath}.bak")
            p.parent.mkdir(parents=True, exist_ok=True)
            p.write_text(content, encoding="utf-8")
            return {"status": "success", "action": "save_file",
                    "path": filepath, "bytes": len(content), "backup_created": backup}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    def delete_file(self, filepath: str, confirm: bool = False) -> Dict[str, Any]:
        """Delete a file or directory. confirm must be True to proceed."""
        if not confirm:
            return {"status": "error",
                    "error": "Set confirm=True to confirm deletion. This is irreversible."}
        filepath = self._host_path(filepath)
        try:
            p = Path(filepath)
            if p.is_dir():
                shutil.rmtree(filepath)
            else:
                p.unlink()
            return {"status": "success", "action": "delete_file", "path": filepath}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    def copy_file(self, src: str, dst: str) -> Dict[str, Any]:
        """Copy a file or directory."""
        src = self._host_path(src)
        dst = self._host_path(dst)
        try:
            Path(dst).parent.mkdir(parents=True, exist_ok=True)
            if Path(src).is_dir():
                shutil.copytree(src, dst, dirs_exist_ok=True)
            else:
                shutil.copy2(src, dst)
            return {"status": "success", "action": "copy_file", "src": src, "dst": dst}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    def move_file(self, src: str, dst: str) -> Dict[str, Any]:
        """Move or rename a file or directory."""
        src = self._host_path(src)
        dst = self._host_path(dst)
        try:
            Path(dst).parent.mkdir(parents=True, exist_ok=True)
            shutil.move(src, dst)
            return {"status": "success", "action": "move_file", "src": src, "dst": dst}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    def list_directory(self, path: str, show_hidden: bool = False,
                       recursive: bool = False) -> Dict[str, Any]:
        """List directory contents with metadata."""
        path = self._host_path(path)
        try:
            entries  = []
            base     = Path(path)
            iterator = base.rglob("*") if recursive else base.iterdir()
            for item in sorted(iterator):
                if not show_hidden and item.name.startswith("."):
                    continue
                stat = item.stat()
                entries.append({
                    "name":       item.name,
                    "path":       str(item),
                    "is_dir":     item.is_dir(),
                    "size_bytes": stat.st_size,
                    "modified":   datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    "extension":  item.suffix.lower(),
                })
            return {"status": "success", "path": path,
                    "entries": entries, "count": len(entries)}
        except Exception as e:
            return {"status": "error", "error": str(e)}


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT EDITING
# ══════════════════════════════════════════════════════════════════════════════

class DocumentTool:
    """
    Read, create, and edit documents.
    Supported formats: .txt .md .json .csv .docx .xlsx
    """

    def read_document(self, filepath: str) -> Dict[str, Any]:
        """Read a document and return structured content."""
        p   = Path(filepath)
        ext = p.suffix.lower()
        try:
            if ext in (".txt", ".md", ".py", ".js", ".ts", ".html", ".css", ".yaml", ".yml"):
                content = p.read_text(encoding="utf-8", errors="replace")
                return {"status": "success", "format": ext, "content": content,
                        "lines": len(content.splitlines())}

            elif ext == ".json":
                data = json.loads(p.read_text(encoding="utf-8"))
                return {"status": "success", "format": "json", "content": data}

            elif ext == ".csv":
                rows = []
                with open(filepath, newline="", encoding="utf-8") as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        rows.append(dict(row))
                return {"status": "success", "format": "csv",
                        "headers": list(rows[0].keys()) if rows else [],
                        "rows": rows, "total_rows": len(rows)}

            elif ext == ".docx":
                err = _require(DOCX_AVAILABLE, "python-docx")
                if err: return err
                doc        = DocxDocument(filepath)
                paragraphs = [p.text for p in doc.paragraphs]
                tables     = [[[ cell.text for cell in row.cells]
                               for row in table.rows]
                              for table in doc.tables]
                return {"status": "success", "format": "docx",
                        "paragraphs": paragraphs, "tables": tables,
                        "paragraph_count": len(paragraphs)}

            elif ext == ".xlsx":
                err = _require(OPENPYXL_AVAILABLE, "openpyxl")
                if err: return err
                wb     = openpyxl.load_workbook(filepath, data_only=True)
                sheets = {name: [list(row) for row in wb[name].iter_rows(values_only=True)]
                          for name in wb.sheetnames}
                return {"status": "success", "format": "xlsx",
                        "sheets": sheets, "sheet_names": wb.sheetnames}

            else:
                content = p.read_text(encoding="utf-8", errors="replace")
                return {"status": "success", "format": "text", "content": content}

        except Exception as e:
            return {"status": "error", "error": str(e)}

    def create_document(self, filepath: str, content: str = "",
                        doc_type: str = "txt") -> Dict[str, Any]:
        """Create a new document. doc_type: txt | md | docx | xlsx | json | csv"""
        try:
            p   = Path(filepath)
            p.parent.mkdir(parents=True, exist_ok=True)
            ext = p.suffix.lower() or f".{doc_type}"

            if ext == ".docx":
                err = _require(DOCX_AVAILABLE, "python-docx")
                if err: return err
                doc = DocxDocument()
                if content:
                    doc.add_paragraph(content)
                doc.save(filepath)
            elif ext == ".xlsx":
                err = _require(OPENPYXL_AVAILABLE, "openpyxl")
                if err: return err
                openpyxl.Workbook().save(filepath)
            else:
                p.write_text(content, encoding="utf-8")

            return {"status": "success", "action": "create_document",
                    "path": filepath, "format": ext}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    def edit_document(self, filepath: str, edits: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Apply structured edits to a document.

        For .txt / .md:
          {"action": "replace",     "find": "old",    "replace_with": "new"}
          {"action": "append",      "content": "..."}
          {"action": "prepend",     "content": "..."}
          {"action": "insert_line", "line_number": 5, "content": "..."}
          {"action": "delete_line", "line_number": 5}

        For .docx:
          {"action": "append_paragraph", "content": "..."}
          {"action": "add_heading",      "content": "...", "level": 1}
          {"action": "add_table",        "rows": 3, "cols": 4}
          {"action": "replace",          "find": "old", "replace_with": "new"}

        For .xlsx:
          {"action": "set_cell",   "sheet": "Sheet1", "cell": "A1", "value": "..."}
          {"action": "append_row", "sheet": "Sheet1", "values": [1, 2, 3]}
          {"action": "add_sheet",  "name": "NewSheet"}
        """
        p       = Path(filepath)
        ext     = p.suffix.lower()
        applied = []
        errors  = []

        try:
            if ext in (".txt", ".md", ".py", ".js", ".html", ".css", ".yaml"):
                content = p.read_text(encoding="utf-8", errors="replace")
                lines   = content.splitlines()
                for edit in edits:
                    action = edit.get("action")
                    try:
                        if action == "replace":
                            content = content.replace(edit["find"], edit["replace_with"])
                            lines   = content.splitlines()
                        elif action == "append":
                            content += "\n" + edit["content"]
                            lines    = content.splitlines()
                        elif action == "prepend":
                            content  = edit["content"] + "\n" + content
                            lines    = content.splitlines()
                        elif action == "insert_line":
                            lines.insert(edit["line_number"] - 1, edit["content"])
                            content = "\n".join(lines)
                        elif action == "delete_line":
                            ln = edit["line_number"] - 1
                            if 0 <= ln < len(lines):
                                lines.pop(ln)
                                content = "\n".join(lines)
                        else:
                            errors.append(f"Unknown action: {action}"); continue
                        applied.append(action)
                    except Exception as e:
                        errors.append(f"{action}: {e}")
                shutil.copy2(filepath, f"{filepath}.bak")
                p.write_text(content, encoding="utf-8")

            elif ext == ".docx":
                err = _require(DOCX_AVAILABLE, "python-docx")
                if err: return err
                doc = DocxDocument(filepath)
                for edit in edits:
                    action = edit.get("action")
                    try:
                        if action == "append_paragraph":
                            doc.add_paragraph(edit["content"])
                        elif action == "add_heading":
                            doc.add_heading(edit["content"], level=edit.get("level", 1))
                        elif action == "add_table":
                            doc.add_table(rows=edit["rows"], cols=edit["cols"])
                        elif action == "replace":
                            for para in doc.paragraphs:
                                for run in para.runs:
                                    run.text = run.text.replace(edit["find"], edit["replace_with"])
                        else:
                            errors.append(f"Unknown action: {action}"); continue
                        applied.append(action)
                    except Exception as e:
                        errors.append(f"{action}: {e}")
                shutil.copy2(filepath, f"{filepath}.bak")
                doc.save(filepath)

            elif ext == ".xlsx":
                err = _require(OPENPYXL_AVAILABLE, "openpyxl")
                if err: return err
                wb = openpyxl.load_workbook(filepath)
                for edit in edits:
                    action     = edit.get("action")
                    sheet_name = edit.get("sheet", wb.sheetnames[0])
                    try:
                        if action == "set_cell":
                            wb[sheet_name][edit["cell"]] = edit["value"]
                        elif action == "append_row":
                            wb[sheet_name].append(edit["values"])
                        elif action == "add_sheet":
                            wb.create_sheet(edit["name"])
                        else:
                            errors.append(f"Unknown action: {action}"); continue
                        applied.append(action)
                    except Exception as e:
                        errors.append(f"{action}: {e}")
                shutil.copy2(filepath, f"{filepath}.bak")
                wb.save(filepath)

            else:
                return {"status": "error", "error": f"Unsupported format: {ext}"}

            return {"status": "success" if not errors else "partial",
                    "path": filepath, "edits_applied": applied,
                    "errors": errors, "backup_created": True}

        except Exception as e:
            return {"status": "error", "error": str(e)}

    def save_document(self, filepath: str, content: Any,
                      backup: bool = True) -> Dict[str, Any]:
        """Save content to document. Handles JSON dicts, CSV lists, and plain text."""
        p   = Path(filepath)
        ext = p.suffix.lower()
        try:
            if backup and p.exists():
                shutil.copy2(filepath, f"{filepath}.bak")
            if ext == ".json":
                p.write_text(json.dumps(content, indent=2, ensure_ascii=False), encoding="utf-8")
            elif ext == ".csv" and isinstance(content, list):
                with open(filepath, "w", newline="", encoding="utf-8") as f:
                    if content and isinstance(content[0], dict):
                        writer = csv.DictWriter(f, fieldnames=content[0].keys())
                        writer.writeheader(); writer.writerows(content)
                    else:
                        csv.writer(f).writerows(content)
            else:
                p.write_text(str(content), encoding="utf-8")
            return {"status": "success", "path": filepath, "backup_created": backup}
        except Exception as e:
            return {"status": "error", "error": str(e)}


# ══════════════════════════════════════════════════════════════════════════════
# BROWSER AUTOMATION (Playwright)
# ══════════════════════════════════════════════════════════════════════════════

class BrowserAutomationTool:
    """Full browser automation via Playwright."""

    def __init__(self):
        self._playwright = None
        self._browser    = None
        self._page       = None
        self._context    = None

    async def _ensure_browser(self, headless: bool = True):
        if not PLAYWRIGHT_AVAILABLE:
            raise RuntimeError("playwright not installed. Run: pip install playwright && playwright install chromium")
        if self._browser is None:
            self._playwright = await async_playwright().start()
            self._browser    = await self._playwright.chromium.launch(
                headless=headless,
                args=["--no-sandbox", "--disable-dev-shm-usage"],
            )
            self._context = await self._browser.new_context(
                viewport={"width": 1280, "height": 800},
                user_agent=(
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/120.0.0.0 Safari/537.36"
                ),
            )
            self._page = await self._context.new_page()

    async def browse_to(self, url: str, wait_until: str = "domcontentloaded",
                        headless: bool = True) -> Dict[str, Any]:
        """Navigate to a URL."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            await self._ensure_browser(headless=headless)
            response = await self._page.goto(url, wait_until=wait_until, timeout=30000)
            return {"status": "success", "url": self._page.url,
                    "title": await self._page.title(),
                    "http_status": response.status if response else None}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_get_text(self, selector: str = "body",
                               limit: int = 5000) -> Dict[str, Any]:
        """Extract visible text from a CSS selector."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            el = await self._page.query_selector(selector)
            if not el:
                return {"status": "error", "error": f"Selector '{selector}' not found"}
            text = await el.inner_text()
            return {"status": "success", "selector": selector,
                    "text": text[:limit], "total_length": len(text)}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_click(self, selector: str) -> Dict[str, Any]:
        """Click an element by CSS selector."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            await self._page.click(selector, timeout=10000)
            return {"status": "success", "action": "click", "selector": selector}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_type(self, selector: str, text: str,
                           clear_first: bool = True) -> Dict[str, Any]:
        """Type text into an input field."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            if clear_first:
                await self._page.fill(selector, "")
            await self._page.type(selector, text, delay=30)
            return {"status": "success", "action": "type", "selector": selector}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_fill_form(self, fields: Dict[str, str],
                                submit_selector: Optional[str] = None) -> Dict[str, Any]:
        """Fill multiple form fields and optionally submit. fields: {selector: value}"""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        filled  = []
        errors_ = []
        try:
            for selector, value in fields.items():
                try:
                    await self._page.fill(selector, value)
                    filled.append(selector)
                except Exception as e:
                    errors_.append(f"{selector}: {e}")
            submitted = False
            if submit_selector:
                try:
                    await self._page.click(submit_selector)
                    submitted = True
                except Exception as e:
                    errors_.append(f"submit: {e}")
            return {"status": "success" if not errors_ else "partial",
                    "filled": filled, "submitted": submitted, "errors": errors_}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_screenshot(self, save_path: str = "/tmp/browser_screenshot.png",
                                  full_page: bool = False) -> Dict[str, Any]:
        """Take a screenshot of the current browser page."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            await self._page.screenshot(path=save_path, full_page=full_page)
            return {"status": "success", "path": save_path,
                    "url": self._page.url, "full_page": full_page}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_scroll(self, direction: str = "down",
                             amount: int = 500) -> Dict[str, Any]:
        """Scroll the page. direction: up | down | top | bottom."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            scripts = {
                "top":    "window.scrollTo(0, 0)",
                "bottom": "window.scrollTo(0, document.body.scrollHeight)",
                "down":   f"window.scrollBy(0, {amount})",
                "up":     f"window.scrollBy(0, -{amount})",
            }
            await self._page.evaluate(scripts.get(direction, f"window.scrollBy(0, {amount})"))
            return {"status": "success", "action": "scroll", "direction": direction}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_back(self) -> Dict[str, Any]:
        """Navigate browser back."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            await self._page.go_back()
            return {"status": "success", "action": "back", "url": self._page.url}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_forward(self) -> Dict[str, Any]:
        """Navigate browser forward."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            await self._page.go_forward()
            return {"status": "success", "action": "forward", "url": self._page.url}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_find_element(self, selector: str) -> Dict[str, Any]:
        """Check if element exists and return its properties."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            el = await self._page.query_selector(selector)
            if not el:
                return {"status": "success", "found": False, "selector": selector}
            return {"status": "success", "found": True, "selector": selector,
                    "text": (await el.inner_text())[:500],
                    "bounding_box": await el.bounding_box(),
                    "is_visible":   await el.is_visible(),
                    "is_enabled":   await el.is_enabled()}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_execute_js(self, script: str) -> Dict[str, Any]:
        """Execute JavaScript in the browser and return the result."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            result = await self._page.evaluate(script)
            return {"status": "success", "result": result}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_get_links(self) -> Dict[str, Any]:
        """Extract all links from the current page."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            links = await self._page.evaluate(
                "() => Array.from(document.querySelectorAll('a[href]')).map(a => "
                "({text: a.innerText.trim().slice(0,100), href: a.href}))"
                ".filter(l => l.href.startsWith('http'))"
            )
            return {"status": "success", "links": links, "count": len(links)}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_download(self, url: str,
                               save_path: str = "/tmp/download") -> Dict[str, Any]:
        """Download a file from URL."""
        err = _require(PLAYWRIGHT_AVAILABLE, "playwright")
        if err: return err
        try:
            async with self._page.expect_download() as dl_info:
                await self._page.evaluate(f"window.location.href = '{url}'")
            dl = await dl_info.value
            await dl.save_as(save_path)
            return {"status": "success", "path": save_path,
                    "suggested_filename": dl.suggested_filename}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    async def browser_close(self) -> Dict[str, Any]:
        """Close the browser and free all resources."""
        try:
            if self._browser:
                await self._browser.close()
                self._browser = self._page = self._context = None
            if self._playwright:
                await self._playwright.stop()
                self._playwright = None
            return {"status": "success", "action": "browser_closed"}
        except Exception as e:
            return {"status": "error", "error": str(e)}

    def get_current_url(self) -> Dict[str, Any]:
        """Return the current browser URL."""
        if self._page is None:
            return {"status": "error", "error": "Browser not open. Call browse_to first."}
        return {"status": "success", "url": self._page.url}


# ══════════════════════════════════════════════════════════════════════════════
# SINGLETON EXPORTS
# ══════════════════════════════════════════════════════════════════════════════

mouse_kb_tool  = MouseKeyboardTool()
file_tool      = FileManagementTool()
document_tool  = DocumentTool()
browser_tool   = BrowserAutomationTool()